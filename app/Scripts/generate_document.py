import sys
import json
import argparse
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def log(msg):
    print(msg, file=sys.stderr)

def is_red_color(rgb_color):
    """Detects standard and custom red shades used for instructions in templates."""
    if not rgb_color:
        return False
    try:
        color_hex = str(rgb_color).upper()
        # Common red hex values in Word templates
        if color_hex in ['FF0000', 'EE0000', 'FF3333', 'ED1C24', 'C00000', '990000', 'E50000', 'FE0000', 'D32F2F', 'F44336']:
            return True
        if len(color_hex) == 6:
            r = int(color_hex[0:2], 16)
            g = int(color_hex[2:4], 16)
            b = int(color_hex[4:6], 16)
            # Red is predominant
            if r > 150 and g < 100 and b < 100:
                return True
    except Exception:
        pass
    return False

def remove_red_text(doc):
    """Wipes out red instructional text from paragraphs and tables."""
    for p in doc.paragraphs:
        for run in p.runs:
            if run.font.color and is_red_color(run.font.color.rgb):
                run.text = ""
                
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        if run.font.color and is_red_color(run.font.color.rgb):
                            run.text = ""

def replace_placeholders_in_paragraph(p, placeholders):
    """Safely replaces placeholders in a paragraph, preserving first run's formatting and avoiding Word run-splitting bugs."""
    if not p.runs:
        return False
    
    p_text = p.text
    has_any = False
    for k in placeholders:
        if k in p_text:
            has_any = True
            break
            
    if not has_any:
        return False
        
    # Replace all matches sequentially
    full_text = p_text
    for k, v in placeholders.items():
        full_text = full_text.replace(k, str(v if v is not None else ""))
        
    # Write the replaced text back to the first run, empty all others
    p.runs[0].text = full_text
    for run in p.runs[1:]:
        run.text = ""
        
    return True

def create_materials_services_table(doc, items):
    """Finds {{ materiais_servicos }} placeholder and inserts a beautiful, formatted table."""
    if not items:
        # Fallback text if no items
        for p in doc.paragraphs:
            if "{{ materiais_servicos }}" in p.text:
                replace_placeholders_in_paragraph(p, {"{{ materiais_servicos }}": "Nenhum material ou serviço cadastrado."})
        return False

    headers = ['ITEM', 'DESCRIÇÃO', 'UNIDADE', 'QTDE.', 'VALOR UNITÁRIO', 'VALOR TOTAL']
    
    # Locate target paragraph containing the placeholder
    target_p = None
    for p in doc.paragraphs:
        if "{{ materiais_servicos }}" in p.text:
            target_p = p
            break
            
    if not target_p:
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if "{{ materiais_servicos }}" in p.text:
                            target_p = p
                            break
                    if target_p: break
                if target_p: break

    if not target_p:
        log("Warning: {{ materiais_servicos }} placeholder not found.")
        return False

    # Create table with 1 header row + items + 1 total row
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    
    # Header styling
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.bold = True
            run.font.name = 'Arial'
            run.font.size = Pt(10)

    total_geral = 0
    for idx, item in enumerate(items):
        qty = float(item.get('quantity', 0))
        val = float(item.get('unit_value', 0))
        total = qty * val
        total_geral += total
        
        row_cells = table.add_row().cells
        
        # Item Index
        row_cells[0].text = str(idx + 1)
        row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Description
        row_cells[1].text = str(item.get('description', ''))
        row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Unit
        row_cells[2].text = str(item.get('unit', ''))
        row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Quantity
        row_cells[3].text = f"{qty:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
        row_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Unit Value
        row_cells[4].text = f"R$ {val:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
        row_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Total Value
        row_cells[5].text = f"R$ {total:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
        row_cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Font styling for item cells
        for cell in row_cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(10)

    # Total Sum Row
    row_cells = table.add_row().cells
    row_cells[0].text = "TOTAL ESTIMADO DA CONTRATAÇÃO:"
    row_cells[0].merge(row_cells[4])
    row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in row_cells[0].paragraphs[0].runs:
        run.font.bold = True
        run.font.name = 'Arial'
        run.font.size = Pt(10)
        
    row_cells[5].text = f"R$ {total_geral:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
    row_cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in row_cells[5].paragraphs[0].runs:
        run.font.bold = True
        run.font.name = 'Arial'
        run.font.size = Pt(10)

    # Place the table right after target placeholder paragraph
    target_p._element.addnext(table._element)
    target_p.text = "" # Clear placeholder text
    return True

def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--template", required=True)
    parser.add_argument("--output", required=True)
    parser.add_argument("--data", required=True)
    args = parser.parse_args()

    try:
        data = json.loads(args.data)
        doc = Document(args.template)
        
        filename = os.path.basename(args.output).upper()
        log(f"Processing document generation for: {filename}")

        # 1. Fill Materials & Services Table
        if "items" in data:
            create_materials_services_table(doc, data["items"])
            
        # 2. Substitute all double-curly bracket placeholders (including signature metadata)
        placeholders = data.get("placeholders", {})
        if placeholders:
            # Substitute in primary paragraphs
            for p in doc.paragraphs:
                replace_placeholders_in_paragraph(p, placeholders)
                
            # Substitute in table paragraphs
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            replace_placeholders_in_paragraph(p, placeholders)

        # 3. Strip red instructional texts cleanly
        remove_red_text(doc)
        
        # 4. Save and export the final high-quality document
        doc.save(args.output)
        log(f"Document generated successfully: {args.output}")
        
    except Exception as e:
        import traceback
        log(traceback.format_exc())
        sys.exit(1)

if __name__ == "__main__":
    main()
