import docx
import os

def inspect_file(filepath):
    print(f"\n=================== INSPECTING: {os.path.basename(filepath)} ===================")
    doc = docx.Document(filepath)
    
    print("\n--- Paragraphs with placeholders or interesting text ---")
    placeholders = []
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if "{{" in text or "}}" in text or "___" in text:
            print(f"Para {i}: {text}")
            placeholders.append(text)
            
    print("\n--- Tables with placeholders or interesting text ---")
    for t_idx, table in enumerate(doc.tables):
        print(f"\nTable {t_idx} (rows: {len(table.rows)}, cols: {len(table.rows[0].cells) if table.rows else 0}):")
        for r_idx, row in enumerate(table.rows):
            row_text = [cell.text.strip() for cell in row.cells]
            if any("{{" in text or "}}" in text or "___" in text or "NOME:" in text.upper() for text in row_text):
                print(f"  Row {r_idx}: {row_text}")

inspect_file("/home/kawan/Documents/code/areas/SECTI/compras-assai/docs/MODELO_SD.docx")
inspect_file("/home/kawan/Documents/code/areas/SECTI/compras-assai/docs/MODELO_ETP.docx")
